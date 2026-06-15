#!/usr/bin/env python3
"""把 Markdown 文件导出为 Word 文档（.docx）。"""

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt


HEADING = re.compile(r'^(#{1,6})\s+(.+?)(?:\s+#+)?$')
CODE_FENCE = re.compile(r'^```')
LIST_ITEM = re.compile(r'^\s*[-*]\s+(.+)$')


def export(md_path: Path, docx_path: Path) -> None:
    """读取 md_path，转换并保存到 docx_path。"""
    doc = Document()
    in_code_block = False

    for raw_line in md_path.read_text(encoding='utf-8').splitlines():
        line = raw_line.rstrip('\n')

        # 处理代码块围栏
        if CODE_FENCE.match(line):
            in_code_block = not in_code_block
            continue

        # 代码块内：等宽字体
        if in_code_block:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(line)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            continue

        # 空行：跳过（python-docx 会自动处理段落间距）
        if not line.strip():
            continue

        # 标题
        m = HEADING.match(line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            doc.add_heading(text, level=min(level, 4))
            continue

        # 列表项
        m = LIST_ITEM.match(line)
        if m:
            doc.add_paragraph(m.group(1).strip(), style='List Bullet')
            continue

        # 普通段落
        doc.add_paragraph(line)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)


def main() -> None:
    parser = argparse.ArgumentParser(
        description='把 Markdown 文件导出为 Word 文档'
    )
    parser.add_argument('input', help='输入的 Markdown 文件路径')
    parser.add_argument(
        '-o', '--output',
        help='输出的 .docx 路径（默认与输入同目录、同文件名）'
    )
    args = parser.parse_args()

    md_path = Path(args.input)
    if not md_path.exists():
        print(f"错误：文件不存在 {md_path}", file=sys.stderr)
        sys.exit(1)

    docx_path = Path(args.output) if args.output else md_path.with_suffix('.docx')

    export(md_path, docx_path)
    print(f"已生成：{docx_path}")


if __name__ == '__main__':
    main()
