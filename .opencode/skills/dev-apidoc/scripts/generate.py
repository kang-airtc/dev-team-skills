#!/usr/bin/env python3
"""dev-apidoc: 从 OpenAPI 规范生成 Word 接口文档"""

import argparse
import json
import sys
from collections import defaultdict
from pathlib import Path
from urllib.request import urlopen

from docx import Document


# 仅展示标准 HTTP 动词
HTTP_METHODS = {'GET', 'POST', 'PUT', 'DELETE', 'PATCH'}


def load_spec(input_path=None, url=None) -> dict:
    """从本地文件或 URL 加载 OpenAPI JSON"""
    if url:
        with urlopen(url, timeout=10) as resp:
            return json.loads(resp.read())
    if input_path:
        return json.loads(Path(input_path).read_text(encoding='utf-8'))
    raise ValueError('必须指定 --input 或 --url')


def collect_endpoints(spec: dict):
    """从 spec 提取所有接口，按 tag 分组。

    返回：[(tag, [(path, method, op), ...]), ...]
    """
    paths = spec.get('paths', {})
    by_tag = defaultdict(list)

    for path, methods in paths.items():
        for method, op in methods.items():
            if method.upper() not in HTTP_METHODS:
                continue
            tags = op.get('tags', ['default']) or ['default']
            for tag in tags:
                by_tag[tag].append((path, method.upper(), op))

    return list(by_tag.items())


def add_overview_table(doc: Document, by_tag: list):
    """添加接口总览表"""
    doc.add_heading('1. 接口总览', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid'
    hdr = table.rows[0].cells
    hdr[0].text = '方法'
    hdr[1].text = '路径'
    hdr[2].text = '摘要'

    for _, endpoints in by_tag:
        for path, method, op in endpoints:
            row = table.add_row().cells
            row[0].text = method
            row[1].text = path
            row[2].text = op.get('summary', '')


def add_parameters(doc: Document, params: list):
    """添加路径/查询参数表"""
    if not params:
        return
    doc.add_heading('参数', level=3)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid'
    hdr = table.rows[0].cells
    hdr[0].text = '名称'
    hdr[1].text = '位置'
    hdr[2].text = '必填'
    hdr[3].text = '说明'
    for p in params:
        row = table.add_row().cells
        row[0].text = p.get('name', '')
        row[1].text = p.get('in', '')
        row[2].text = '是' if p.get('required') else '否'
        row[3].text = p.get('description', '')


def add_request_body(doc: Document, req_body: dict):
    """添加请求体说明"""
    if not req_body:
        return
    doc.add_heading('请求体', level=3)
    for mime, schema_info in req_body.get('content', {}).items():
        doc.add_paragraph(f'Content-Type: {mime}')
        schema = schema_info.get('schema', {})
        if '$ref' in schema:
            ref_name = schema['$ref'].split('/')[-1]
            doc.add_paragraph(f'Schema: {ref_name}')
        elif 'type' in schema:
            doc.add_paragraph(f"Type: {schema['type']}")


def add_responses(doc: Document, responses: dict):
    """添加响应表（按统一 {code, msg, data} 格式呈现）"""
    if not responses:
        return
    doc.add_heading('响应', level=3)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid'
    hdr = table.rows[0].cells
    hdr[0].text = '状态码'
    hdr[1].text = '说明'
    for status, resp in responses.items():
        row = table.add_row().cells
        row[0].text = status
        row[1].text = resp.get('description', '')


def add_endpoint(doc: Document, path: str, method: str, op: dict):
    """添加单个接口的详细描述"""
    doc.add_heading(f'{method} {path}', level=2)
    if op.get('summary'):
        doc.add_paragraph(op['summary'])
    if op.get('description'):
        doc.add_paragraph(op['description'])

    add_parameters(doc, op.get('parameters', []))
    add_request_body(doc, op.get('requestBody', {}))
    add_responses(doc, op.get('responses', {}))


def add_error_codes_section(doc: Document):
    """所有接口共用的错误码段说明"""
    doc.add_heading(f'{99}. 通用错误码', level=1)
    doc.add_paragraph(
        '所有接口遵循统一的 {code, msg, data} 响应格式，错误码段约定如下：'
    )
    doc.add_paragraph('0：成功', style='List Bullet')
    doc.add_paragraph('1000-1999：通用错误（参数、鉴权、限流）', style='List Bullet')
    doc.add_paragraph('2000-2999：业务错误', style='List Bullet')
    doc.add_paragraph(
        '详细错误码字典见 dev-backend-lint/references/error-codes.md。'
    )


def generate_doc(spec: dict, output_path: Path):
    """主流程"""
    doc = Document()

    # 标题
    info = spec.get('info', {})
    title = info.get('title', 'API Documentation')
    version = info.get('version', '1.0')
    doc.add_heading(f'{title}（v{version}）', level=0)

    if info.get('description'):
        doc.add_paragraph(info['description'])

    # 收集并分组接口
    by_tag = collect_endpoints(spec)

    # 1. 总览表
    add_overview_table(doc, by_tag)

    # 2~N. 按 tag 分组的详细接口
    for idx, (tag, endpoints) in enumerate(by_tag, start=2):
        doc.add_heading(f'{idx}. {tag}', level=1)
        for path, method, op in endpoints:
            add_endpoint(doc, path, method, op)

    # 末尾：通用错误码
    add_error_codes_section(doc)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser(description='OpenAPI → Word 接口文档')
    parser.add_argument('--url', '-u',
                        help='OpenAPI JSON 端点 URL（如 http://localhost:8000/openapi.json）')
    parser.add_argument('--input', '-i', help='本地 OpenAPI JSON 文件路径')
    parser.add_argument('--output', '-o', required=True, help='输出 .docx 路径')
    args = parser.parse_args()

    if not args.url and not args.input:
        print('错误：必须指定 --url 或 --input', file=sys.stderr)
        sys.exit(1)

    try:
        spec = load_spec(args.input, args.url)
    except Exception as e:
        print(f'错误：加载 OpenAPI 失败 - {e}', file=sys.stderr)
        sys.exit(1)

    output_path = Path(args.output)
    generate_doc(spec, output_path)
    print(f'已生成：{output_path}')


if __name__ == '__main__':
    main()
